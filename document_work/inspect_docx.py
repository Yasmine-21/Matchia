from pathlib import Path
from docx import Document

for path in [
    Path(r"D:\PFE M2\Platforme SaaS\document_work\Assortis-ICA-IBF.docx"),
    Path(r"D:\PFE M2\Platforme SaaS\document_work\Developpements_Assortis_ICA_IBF_verifie.docx"),
]:
    document = Document(path)
    print(f"--- {path.name} ---")
    print(f"PARAGRAPHS={len(document.paragraphs)} TABLES={len(document.tables)}")
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            print(f"P{index}: {text}")
    for table_index, table in enumerate(document.tables, start=1):
        print(f"TABLE {table_index}: {len(table.rows)}x{len(table.columns)}")
        for row in table.rows:
            print(" | ".join(cell.text.replace("\n", " / ") for cell in row.cells))
