from docx import Document
from pathlib import Path

path = Path(r"D:\PFE M2\Platforme SaaS\Developpements_Assortis_ICA_IBF_avec_captures.docx")
document = Document(path)
text = "\n".join(p.text for p in document.paragraphs)
for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            text += "\n" + cell.text

print(f"tables={len(document.tables)}")
print(f"inline_shapes={len(document.inline_shapes)}")
print(f"old_placeholder_occurrences={text.count('Capture non insérée')}")
print(f"scope_notice_present={'Mise à jour des preuves visuelles' in text}")
print(f"file_size={path.stat().st_size}")
