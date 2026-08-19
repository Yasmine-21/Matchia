from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


WORK = Path(r"D:\PFE M2\Platforme SaaS")
SOURCE = WORK / "document_work" / "Developpements_Assortis_ICA_IBF_verifie.docx"
OUTPUT = WORK / "Developpements_Assortis_ICA_IBF_avec_captures.docx"
SCREENSHOT = WORK / "evidence_screenshots" / "organization-user-dashboard.png"

document = Document(SOURCE)

# Insert a concise audit-scope notice after the existing introduction.  The
# rest of the source report remains intact, preserving its existing structure.
anchor = document.paragraphs[5]
notice = document.add_paragraph()
notice.style = anchor.style
notice.alignment = WD_ALIGN_PARAGRAPH.LEFT
notice.paragraph_format.space_before = Pt(6)
notice.paragraph_format.space_after = Pt(8)
run = notice.add_run(
    "Mise à jour des preuves visuelles - 18 août 2026. La capture insérée ci-dessous "
    "a été obtenue sur l’espace connecté organization-user. La session existante a permis "
    "de vérifier le tableau de bord et l’accès aux modules principaux. Après déconnexion, "
    "les tentatives de connexion aux comptes admin, expert et organization sont restées sur "
    "le formulaire de connexion, sans message d’erreur ni ouverture de session. Les éléments "
    "qui dépendent de ces rôles restent donc classés Non vérifiable. Les captures ne prouvent "
    "que les écrans et modules effectivement visibles ; elles ne prouvent pas les flux non exécutés."
)
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(89, 89, 89)
anchor._p.addnext(notice._p)

# Screenshot placeholders follow each functional assessment.  Reuse the
# authenticated dashboard only where it visibly exposes the named module.
dashboard_evidence = {
    4: "Preuve visuelle - Matching Projects et ses alertes sont proposés depuis le tableau de bord.",
    6: "Preuve visuelle - l’accès Experts est proposé depuis le tableau de bord ; les flux experts détaillés ne sont pas démontrés.",
    14: "Preuve visuelle - Matching Projects est présenté comme un module de matching personnalisé.",
    16: "Preuve visuelle - l’accès My Organization est visible depuis le tableau de bord.",
    30: "Preuve visuelle - l’accès à l’espace organisation est visible depuis le tableau de bord.",
    32: "Preuve visuelle - le module Training est visible depuis le tableau de bord.",
    46: "Preuve visuelle - Matching Projects et My Organization sont visibles depuis le tableau de bord.",
    48: "Preuve visuelle - le Posting Board est décrit comme un espace de publication et de gestion des vacances.",
    52: "Preuve visuelle - le tableau de bord expose Matching Projects et ses alertes personnalisées.",
    54: "Preuve visuelle - le tableau de bord expose l’accès My Projects.",
    56: "Preuve visuelle - le tableau de bord expose My Projects et l’accès à la recherche d’experts.",
    58: "Preuve visuelle - l’accès Experts est visible ; la recherche avancée détaillée n’est pas prouvée par cette capture.",
    62: "Preuve visuelle - le Posting Board est accessible depuis le tableau de bord.",
    64: "Preuve visuelle - l’accès My Organization et My Account est visible depuis le tableau de bord.",
    66: "Preuve visuelle - l’accès Statistics est visible depuis le tableau de bord.",
}

for table_number in range(4, len(document.tables) + 1, 2):
    table = document.tables[table_number - 1]
    if len(table.rows) != 1 or len(table.columns) != 1:
        continue
    cell = table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if table_number in dashboard_evidence:
        caption = paragraph.add_run(dashboard_evidence[table_number])
        caption.bold = True
        caption.font.size = Pt(8)
        caption.font.color.rgb = RGBColor(68, 68, 68)
        paragraph.add_run("\n")
        picture_run = paragraph.add_run()
        picture_run.add_picture(str(SCREENSHOT), width=Inches(5.85))
        note = cell.add_paragraph("Capture : tableau de bord authentifié organization-user (18 août 2026).")
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note.runs[0].italic = True
        note.runs[0].font.size = Pt(7.5)
    else:
        note = paragraph.add_run(
            "Aucune capture insérée : la fonctionnalité correspondante n’était pas visible "
            "dans le périmètre effectivement accessible, ou son flux détaillé n’a pas été exécuté."
        )
        note.italic = True
        note.font.size = Pt(8)
        note.font.color.rgb = RGBColor(89, 89, 89)

# Make statuses clearer without changing their factual wording.
for table_number in range(3, len(document.tables) + 1, 2):
    table = document.tables[table_number - 1]
    for row in table.rows:
        if len(row.cells) != 2:
            continue
        if row.cells[0].text.strip().lower() == "statut":
            value = row.cells[1].text.strip().lower()
            color = RGBColor(89, 89, 89)
            if "partiellement" in value:
                color = RGBColor(191, 115, 0)
            elif "implémenté" in value:
                color = RGBColor(46, 125, 50)
            for p in row.cells[1].paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = color

document.save(OUTPUT)
print(OUTPUT)
